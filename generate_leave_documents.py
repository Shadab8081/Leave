#!/usr/bin/env python3
"""
Automated Leave Application & Clearance Generator
====================================================
Reads one Excel input sheet of employees and auto-generates, per employee:
  - Leave Application Form  (common template, sponsor heading swapped)
  - Clearance Form          (sponsor-specific template)
  - Al Falak employees also get the "Leave and Visa Request Form" (PDF)
  - Mehan employees also get the "Vacation Request Format" (docx)

Usage:
    python3 generate_leave_documents.py INPUT.xlsx [--outdir "Leave Documents"]

Re-usable: keep using the same column headers in your Excel file and just
replace/add employee rows, then re-run this script.
"""

import sys
import os
import re
import shutil
import argparse
import subprocess
import datetime as dt
from pathlib import Path

import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATES = SCRIPT_DIR / "templates"
SCRIPTS = SCRIPT_DIR / "scripts"
SOFFICE = SCRIPTS / "office" / "soffice.py"
RECALC = SCRIPTS / "recalc.py"

# ---------------------------------------------------------------------------
# Column mapping — accepts either the user's real sheet headers or the
# original spec's more verbose headers. Add synonyms here if your sheet
# uses different wording.
# ---------------------------------------------------------------------------
COLUMN_SYNONYMS = {
    "employee_id":      ["EMP ID", "EMPLOYEE ID", "EMP. ID"],
    "employee_name":    ["EMP NAME", "EMPLOYEE NAME"],
    "position":         ["JOB TITTLE", "JOB TITLE", "POSITION"],
    "nationality":      ["NATIONALITY"],
    "sponsor_id":       ["SPONSOR ID", "ID NUMBER", "ID#", "ID #"],
    "iqama_number":     ["IQAMA NUMBER", "IQAMA NO", "IQAMA #"],
    "iqama_exp":        ["IQAMA EXP", "IQAMA EXPIRY", "IQAMA EXPIRY DATE"],
    "leave_type":       ["TYPE OF VACATION", "LEAVE TYPE"],
    "leave_start":      ["VAC. START DATE", "VAC START DATE", "LEAVE START DATE"],
    "leave_end":        ["VAC. END DATE", "VAC END DATE", "LEAVE END DATE"],
    "num_days":         ["NO OF DAYS", "NO. OF DAYS", "NUMBER OF LEAVE DAYS"],
    "project":          ["PROJECT"],
    "phone":            ["PHONE NUMBER", "CONTACT NUMBER"],
    "hiring_date":      ["HIRING DATE"],
    "effective_date":   ["EFECTIVE DATE", "EFFECTIVE DATE"],
    "department":       ["DEPARTMENT"],
    "location":         ["LOCATION"],
    "contract":         ["CONTRACT"],
    "entitlement":      ["ENTITLEMENT"],
    "employment_status": ["EMPLOPYEMENT STATUS", "EMPLOYMENT STATUS"],
    "paid_days_leave":  ["PAID DAY LEAVE", "PAID DAYS LEAVE"],
    "unpaid_days_leave": ["UNPAID DAY LEAVE", "UNPAID DAYS LEAVE"],
    "air_ticket":       ["AIR TICKET"],
    "ticket_type":      ["TICKET TYPE"],
    "ticket_class":     ["TICKET CLASS"],
    "exit_reentry_visa": ["EXIT & RE-ENTRY VISA", "EXIT AND RE-ENTRY VISA"],
    "visa_type":        ["VISA TYPE"],
    "visa_days":        ["VISA DAYS"],
    "destination_country": ["DESTINATION COUNTRY"],
    "departure_date":  ["DEPARTURE DATE"],
    "arrival_date":    ["ARRIVAL DATE"],
}

REQUIRED_FIELDS = [
    "employee_name", "sponsor_id", "leave_type", "leave_start", "leave_end",
]

SPONSORS = {
    "FMCO": {
        "match": lambda sid: sid.upper().startswith("7"),
        "heading": "FMCO",
        "folder": "FMCO",
        "clearance_template": TEMPLATES / "fmco_clearance.docx",
        "extra_docs": [],
    },
    "MEHAN": {
        "match": lambda sid: sid.upper().startswith("C"),
        "heading": "MEHAN",
        "folder": "Mehan",
        "clearance_template": TEMPLATES / "mehan_clearance.docx",
        "extra_docs": ["mehan_vacation"],
    },
    "ALFALAK": {
        "match": lambda sid: sid.upper().startswith("EM-"),
        "heading": "AL FALAK",
        "folder": "Al Falak",
        "clearance_template": TEMPLATES / "alfalak_clearance.docx",
        "extra_docs": ["alfalak_pdf"],
    },
}

LEAVE_TYPES = {
    "ANNUAL": "ANNUAL LEAVE",
    "PARTIAL": "PARTIAL LEAVE",
    "EMERGENCY": "EMERGENCY LEAVE",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def normalize_header(h):
    if h is None:
        return ""
    h = str(h).replace("\xa0", " ")
    h = re.sub(r"\s+", " ", h).strip().upper()
    return h


def build_column_map(header_row):
    """Map normalized header text -> our internal field name."""
    lookup = {}
    for field, synonyms in COLUMN_SYNONYMS.items():
        for syn in synonyms:
            lookup[normalize_header(syn)] = field

    col_map = {}
    for idx, cell_val in enumerate(header_row, start=1):
        norm = normalize_header(cell_val)
        if norm in lookup:
            col_map[lookup[norm]] = idx
    return col_map


def to_date(value):
    if value is None or value == "":
        return None
    if isinstance(value, dt.datetime):
        return value.date()
    if isinstance(value, dt.date):
        return value
    s = str(value).strip()
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d %B %Y", "%d %b %Y", "%m/%d/%Y"):
        try:
            return dt.datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def fmt_date(d):
    if d is None:
        return ""
    return d.strftime("%d %B %Y").upper()


def detect_sponsor(sponsor_id):
    if not sponsor_id:
        return None
    sid = str(sponsor_id).strip()
    for name, cfg in SPONSORS.items():
        if cfg["match"](sid):
            return name
    return None


def safe_filename(s):
    s = re.sub(r"[^A-Za-z0-9 _-]", "", str(s))
    return re.sub(r"\s+", "", s).strip("_")


def run_soffice_to_pdf(src_path: Path, outdir: Path):
    outdir.mkdir(parents=True, exist_ok=True)
    cmd = [sys.executable, str(SOFFICE), "--headless", "--convert-to", "pdf",
           "--outdir", str(outdir), str(src_path)]
    subprocess.run(cmd, check=True, capture_output=True)
    pdf_path = outdir / (src_path.stem + ".pdf")
    return pdf_path


# ---------------------------------------------------------------------------
# Reading input
# ---------------------------------------------------------------------------
class Employee:
    def __init__(self, row_num, data):
        self.row_num = row_num
        self.data = data
        self.errors = []

    def get(self, field, default=""):
        return self.data.get(field, default)


def read_employees(input_path):
    wb = openpyxl.load_workbook(input_path, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("Input Excel sheet is empty.")

    header_row = rows[0]
    col_map = build_column_map(header_row)

    missing_cols = [f for f in REQUIRED_FIELDS if f not in col_map]
    if missing_cols:
        raise ValueError(
            "Could not find required column(s) in the input sheet: "
            f"{missing_cols}. Found headers: {list(header_row)}"
        )

    employees = []
    for i, row in enumerate(rows[1:], start=2):
        if all(v is None or str(v).strip() == "" for v in row):
            continue
        data = {}
        for field, col_idx in col_map.items():
            val = row[col_idx - 1] if col_idx - 1 < len(row) else None
            data[field] = val
        employees.append(Employee(i, data))
    return employees


def validate_employee(emp):
    errors = []
    name = emp.get("employee_name")
    if not name or not str(name).strip():
        errors.append("Employee Name is missing")

    sponsor_id = emp.get("sponsor_id")
    if not sponsor_id or not str(sponsor_id).strip():
        errors.append("Sponsor ID (ID NUMBER) is missing")
    else:
        sponsor = detect_sponsor(sponsor_id)
        if sponsor is None:
            errors.append(f"Invalid Sponsor ID format: {sponsor_id!r}")

    leave_type_raw = str(emp.get("leave_type") or "").strip().upper()
    if not leave_type_raw:
        errors.append("Leave Type is missing")
    else:
        if not any(k in leave_type_raw for k in LEAVE_TYPES.keys()):
            errors.append(f"Leave Type not recognized: {leave_type_raw!r}")

    start = to_date(emp.get("leave_start"))
    end = to_date(emp.get("leave_end"))
    if not start:
        errors.append("Leave Start Date is missing or invalid")
    if not end:
        errors.append("Leave End Date is missing or invalid")
    if start and end and end < start:
        errors.append("Leave End Date is before Leave Start Date")

    emp.errors = errors
    return errors


def normalize_leave_type(raw):
    raw_u = str(raw or "").strip().upper()
    for key, label in LEAVE_TYPES.items():
        if key in raw_u:
            return key, label
    return None, raw_u


def visa_days_from_leave(num_days):
    """Auto-select visa validity tier based on total leave days taken."""
    try:
        n = float(str(num_days).strip())
    except (TypeError, ValueError):
        return "60 DAYS"
    if n <= 30:
        return "30 DAYS"
    if n <= 60:
        return "60 DAYS"
    return "90 DAYS"  # cap at 90; longer leave should really use a Multiple-entry visa


# ---------------------------------------------------------------------------
# Document 1: Common Leave Application Form (xlsx -> pdf)
# ---------------------------------------------------------------------------
LEAVE_TYPE_MARK_ROW = {"ANNUAL": 13, "PARTIAL": 14, "EMERGENCY": 16}  # rows in U col


def fill_leave_application(emp, sponsor_name, workdir: Path) -> Path:
    wb = openpyxl.load_workbook(TEMPLATES / "common_leave_form.xlsx")
    ws = wb["FMCO"]  # sheet name in template; content is generic

    heading = SPONSORS[sponsor_name]["heading"]
    ws["G1"] = f"LEAVE APPLICATION FORM ({heading})"

    today = dt.date.today()
    ws["P3"] = dt.datetime.combine(today, dt.time())
    ws["Q5"] = str(emp.get("sponsor_id") or "")

    ws["E9"] = str(emp.get("employee_name") or "")
    ws["E10"] = str(emp.get("position") or "")
    ws["E11"] = str(emp.get("nationality") or "")
    ws["E12"] = str(emp.get("project") or "")
    ws["E13"] = str(emp.get("phone") or "")

    hiring = to_date(emp.get("hiring_date"))
    if hiring:
        ws["O9"] = dt.datetime.combine(hiring, dt.time())

    start = to_date(emp.get("leave_start"))
    end = to_date(emp.get("leave_end"))
    if start:
        ws["E19"] = dt.datetime.combine(start, dt.time())
        last_working_day = start - dt.timedelta(days=1)
        ws["E17"] = dt.datetime.combine(last_working_day, dt.time())
    if end:
        ws["O19"] = dt.datetime.combine(end, dt.time())
        resume = end + dt.timedelta(days=1)
        ws["O17"] = dt.datetime.combine(resume, dt.time())

    lt_key, lt_label = normalize_leave_type(emp.get("leave_type"))
    ws["O13"] = lt_label
    if lt_key and lt_key in LEAVE_TYPE_MARK_ROW:
        ws[f"T{LEAVE_TYPE_MARK_ROW[lt_key]}"] = "X"

    # Fill fields we have data for; leave blank only what's truly absent.
    if emp.get("iqama_number"):
        ws["O10"] = str(emp.get("iqama_number"))
    else:
        ws["O10"] = None
    ws["O11"] = str(emp.get("contract") or "") or None
    ws["O12"] = str(emp.get("entitlement") or "") or None
    ws["E21"] = str(emp.get("employment_status") or "") or None
    ws["K21"] = str(emp.get("paid_days_leave") or "") or None
    ws["Q21"] = str(emp.get("unpaid_days_leave") or "") or None
    ws["D25"] = str(emp.get("air_ticket") or "") or None
    ws["J25"] = str(emp.get("ticket_type") or "ROUNDTRIP")
    ws["P25"] = str(emp.get("ticket_class") or "ECONOMY CLASS")
    ws["F27"] = str(emp.get("exit_reentry_visa") or "") or None
    ws["K27"] = str(emp.get("visa_type") or "SINGLE VISA")
    ws["P27"] = str(emp.get("visa_days") or visa_days_from_leave(emp.get("num_days")))
    ws["F29"] = str(emp.get("destination_country") or "") or None

    departure = to_date(emp.get("departure_date"))
    ws["F31"] = dt.datetime.combine(departure, dt.time()) if departure else None
    arrival = to_date(emp.get("arrival_date"))
    ws["O31"] = dt.datetime.combine(arrival, dt.time()) if arrival else None

    ws["B46"] = str(emp.get("employee_name") or "")
    ws["B47"] = str(emp.get("position") or "")
    for cell in ("C45", "H45", "L45", "P45"):
        ws[cell] = dt.datetime.combine(today, dt.time())

    out_xlsx = workdir / "leave_application.xlsx"
    wb.save(out_xlsx)

    subprocess.run([sys.executable, str(RECALC), str(out_xlsx)],
                    check=False, capture_output=True)

    pdf_path = run_soffice_to_pdf(out_xlsx, workdir)
    return pdf_path


# ---------------------------------------------------------------------------
# Document 2: Clearance Form (docx -> pdf), one per sponsor
# ---------------------------------------------------------------------------
def set_cell_value(cell, text):
    text = "" if text is None else str(text)
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

    # Some cells (e.g. Mehan's "Clearance For") also contain a Word dropdown
    # content control (<w:sdt>) with its own leftover text baked into the
    # template, sitting right next to the plain run above — left alone it
    # renders as duplicated text ("EMERGENCY LEAVEEMERGENCY LEAVE"). Blank
    # those out so only our plain-run value shows.
    for sdt in cell._tc.iter(qn("w:sdt")):
        for t in sdt.iter(qn("w:t")):
            t.text = ""


CLEARANCE_LABELS = {
    "employee name": "employee_name",
    "employee id #": "sponsor_id",
    "hiring date": "hiring_date",
    "iqama / natl' id": "iqama_number",
    "iqama / natl\u2019 id": "iqama_number",
    "position": "position",
    "project": "project",
    "department": "department",
    "location": "location",
    "effective date": "effective_date",
    "clearance for": "clearance_for",
}


def unique_row_cells(row):
    """python-docx repeats the same cell object across a horizontal merge
    (gridSpan) — dedupe by the underlying XML element so we get one entry
    per visually distinct cell, in left-to-right order. Table layouts vary
    between sponsors (6 vs 8 raw grid columns), so this is what lets the
    same fill logic work for all of them."""
    seen = set()
    uniq = []
    for c in row.cells:
        key = id(c._tc)
        if key not in seen:
            seen.add(key)
            uniq.append(c)
    return uniq


def box_cell(cell):
    """Give a cell a full single-line border on all sides — used for the
    'Clearance For' value cell, whose template has top/right/bottom
    explicitly set to no border, leaving the text floating outside any box."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        from docx.oxml import OxmlElement
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            from docx.oxml import OxmlElement
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def fill_clearance_form(emp, sponsor_name, lt_label, workdir: Path) -> Path:
    template = SPONSORS[sponsor_name]["clearance_template"]
    doc = Document(template)

    values = {
        "employee_name": emp.get("employee_name"),
        "sponsor_id": emp.get("sponsor_id"),
        "hiring_date": fmt_date(to_date(emp.get("hiring_date"))),
        "iqama_number": emp.get("iqama_number") or "",
        "position": emp.get("position"),
        "project": emp.get("project"),
        "department": emp.get("department") or emp.get("project") or "",
        "location": emp.get("location") or "",
        "effective_date": fmt_date(to_date(emp.get("effective_date"))),
        "clearance_for": lt_label,
    }

    # Employee-details fields live in the first table for every sponsor.
    # "Clearance For" is in that same table for FMCO/Al Falak, but Mehan
    # keeps it in a separate small table — so that one label is checked
    # everywhere, while the rest are scoped to table[0] only. Scanning every
    # table for all labels would also match unrelated cells like the
    # "PROJECT: Tools" row inside the department checklist table below.
    tables = doc.tables
    if tables:
        for row in tables[0].rows:
            uniq = unique_row_cells(row)
            for i, cell in enumerate(uniq):
                label = normalize_header(cell.text).lower()
                if label in CLEARANCE_LABELS and i + 1 < len(uniq):
                    field = CLEARANCE_LABELS[label]
                    set_cell_value(uniq[i + 1], values.get(field, ""))
                    if field == "clearance_for":
                        box_cell(uniq[i + 1])

    for table in tables[1:]:
        for row in table.rows:
            uniq = unique_row_cells(row)
            for i, cell in enumerate(uniq):
                label = normalize_header(cell.text).lower()
                if label == "clearance for" and i + 1 < len(uniq):
                    set_cell_value(uniq[i + 1], values["clearance_for"])
                    box_cell(uniq[i + 1])

    out_docx = workdir / "clearance_form.docx"
    doc.save(out_docx)
    pdf_path = run_soffice_to_pdf(out_docx, workdir)
    return pdf_path


# ---------------------------------------------------------------------------
# Extra doc: Mehan Vacation Request Format (docx -> pdf)
# ---------------------------------------------------------------------------
def fill_mehan_vacation(emp, lt_key, lt_label, workdir: Path) -> Path:
    doc = Document(TEMPLATES / "mehan_vacation.docx")
    t0 = doc.tables[0]
    # row1: Date: | value | Contract #: | value
    set_cell_value(t0.rows[1].cells[1], fmt_date(dt.date.today()))
    set_cell_value(t0.rows[1].cells[3], str(emp.get("contract") or "2 YEARS"))
    # row2: Employee name | value | Employee MEHAN ID | value
    set_cell_value(t0.rows[2].cells[1], emp.get("employee_name"))
    set_cell_value(t0.rows[2].cells[3], emp.get("sponsor_id"))
    # row3: Iqama # | value | Contract Start Date: | value
    set_cell_value(t0.rows[3].cells[1], emp.get("iqama_number") or "")
    set_cell_value(t0.rows[3].cells[3], fmt_date(to_date(emp.get("hiring_date"))))

    t1 = doc.tables[1]

    # Vacation Type checkboxes: unique cells = [label, [X]Annual, 'Annual/Paid',
    # [X]Emergency, 'Emergency', [X]Non-Paid, 'Non-Paid']
    vac_type_row = unique_row_cells(t1.rows[0])
    checkbox_by_leave = {"ANNUAL": 1, "EMERGENCY": 3, "PARTIAL": 5}
    box_idx = checkbox_by_leave.get(lt_key, 5)  # default to Non-Paid if unrecognized
    if box_idx < len(vac_type_row):
        set_cell_value(vac_type_row[box_idx], "X")

    set_cell_value(t1.rows[1].cells[1], lt_label)
    start = to_date(emp.get("leave_start"))
    end = to_date(emp.get("leave_end"))
    set_cell_value(t1.rows[2].cells[2], fmt_date(start) if start else "")
    set_cell_value(t1.rows[2].cells[5], fmt_date(end) if end else "")
    set_cell_value(t1.rows[2].cells[8], str(emp.get("num_days") or ""))

    out_docx = workdir / "mehan_vacation_form.docx"
    doc.save(out_docx)
    pdf_path = run_soffice_to_pdf(out_docx, workdir)
    return pdf_path


# ---------------------------------------------------------------------------
# Extra doc: Al Falak Leave and Visa Request Form (flat PDF -> overlay -> pdf)
# ---------------------------------------------------------------------------
def fmt_date_short(d):
    """Compact date format for tight PDF form fields, e.g. '11 SEP 2026'."""
    if d is None:
        return ""
    return d.strftime("%d %b %Y").upper()


def visa_validity_checkbox(num_days):
    """Map total leave days to the nearest visa-validity checkbox on the
    Al Falak form: (x-position, label) for the 1/2/3/4/6 Month / Multiple
    checkboxes."""
    try:
        n = float(str(num_days).strip())
    except (TypeError, ValueError):
        n = 0
    tiers = [
        (30, 168.7, "1 Month"),
        (60, 228.4, "2 Months"),
        (90, 298.2, "3 Months"),
        (120, 362.0, "4 Months"),
        (180, 418.0, "6 Months"),
    ]
    for limit, x, label in tiers:
        if n <= limit:
            return x, label
    return 506.3, "Multiple"


def fill_alfalak_pdf(emp, lt_key, workdir: Path) -> Path:
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader, PdfWriter

    page_w, page_h = 612, 792

    def y(top):
        return page_h - top

    overlay_path = workdir / "_alfalak_overlay.pdf"
    c = canvas.Canvas(str(overlay_path), pagesize=(page_w, page_h))
    c.setFont("Helvetica", 9)

    name = str(emp.get("employee_name") or "")
    c.drawString(185, y(125), name)
    c.drawString(462, y(125), str(emp.get("sponsor_id") or ""))
    c.drawString(170, y(137), str(emp.get("project") or ""))

    # Type of Leave checkbox
    checkbox_x = {"ANNUAL": 337, "EMERGENCY": 455}
    if lt_key in checkbox_x:
        c.setFont("Helvetica-Bold", 9)
        c.drawString(checkbox_x[lt_key], y(153.5), "X")
        c.setFont("Helvetica", 9)
    else:
        # No dedicated checkbox for Partial Leave on this template.
        c.setFont("Helvetica", 7)
        c.drawString(500, y(158), "(Partial)")
        c.setFont("Helvetica", 9)

    start = to_date(emp.get("leave_start"))
    end = to_date(emp.get("leave_end"))
    num_days = emp.get("num_days") or ""

    # FROM / TO / TOTAL DAYS — columns are FROM:128-247, TO:247-375,
    # TOTAL DAYS:376-594 (wide). Value row sits right under the headers.
    c.setFont("Helvetica", 8)
    c.drawString(135, y(201), fmt_date(start) if start else "")
    c.drawString(255, y(201), fmt_date(end) if end else "")
    c.drawString(460, y(201), str(num_days))
    c.setFont("Helvetica", 9)

    # VISA NEEDED — always tick YES per company policy
    c.setFont("Helvetica-Bold", 9)
    c.drawString(149, y(224.5), "X")
    c.setFont("Helvetica", 9)

    # WITHOUT PAY — NO. OF DAYS: unpaid leave days
    unpaid = emp.get("unpaid_days_leave") or ""
    if unpaid:
        c.setFont("Helvetica", 8)
        c.drawString(432, y(230.5), str(unpaid))
        c.setFont("Helvetica", 9)

    # VISA VALIDITY — tick the tier matching total leave days
    vx, _ = visa_validity_checkbox(num_days)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(vx, y(258.5), "X")
    c.setFont("Helvetica", 9)

    # IQAMA No. / Iqamma Expiry Date — narrow columns, compact font/format
    c.setFont("Helvetica", 8)
    c.drawString(125, y(316), str(emp.get("iqama_number") or ""))
    iqama_exp = to_date(emp.get("iqama_exp"))
    c.drawString(310, y(316), fmt_date_short(iqama_exp) if iqama_exp else "")
    c.setFont("Helvetica", 9)

    # LAST DAY OF WORK / DATE OF RETURN TO WORK — same line as label,
    # narrow columns so use a compact numeric date.
    last_day = (start - dt.timedelta(days=1)) if start else None
    resume = (end + dt.timedelta(days=1)) if end else None
    c.setFont("Helvetica", 7)
    c.drawString(149, y(548), last_day.strftime("%d/%m/%y") if last_day else "")
    c.drawString(440, y(548), resume.strftime("%d/%m/%y") if resume else "")
    c.setFont("Helvetica", 9)

    c.save()

    base_reader = PdfReader(str(TEMPLATES / "alfalak_leave_visa.pdf"))
    overlay_reader = PdfReader(str(overlay_path))
    writer = PdfWriter()
    base_page = base_reader.pages[0]
    base_page.merge_page(overlay_reader.pages[0])
    writer.add_page(base_page)

    out_pdf = workdir / "alfalak_leave_visa_form.pdf"
    with open(out_pdf, "wb") as f:
        writer.write(f)
    overlay_path.unlink(missing_ok=True)
    return out_pdf


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------
def run_generation(input_path, outdir, log=print):
    """Core pipeline, callable from the CLI or the GUI.

    log: a callable taking one string argument, used for progress messages.
    Returns (generated, error_report) lists.
    """
    input_path = Path(input_path)
    outdir = Path(outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    employees = read_employees(input_path)

    error_report = []
    generated = []

    log(f"Found {len(employees)} employee row(s) in the input sheet.")

    for emp in employees:
        errors = validate_employee(emp)
        if errors:
            for e in errors:
                error_report.append(
                    f"Row {emp.row_num} - {emp.get('employee_name') or '(no name)'}\n"
                    f"    Missing/Invalid: {e}"
                )
            continue

        sponsor_name = detect_sponsor(emp.get("sponsor_id"))
        sponsor_cfg = SPONSORS[sponsor_name]
        lt_key, lt_label = normalize_leave_type(emp.get("leave_type"))

        emp_name_safe = safe_filename(emp.get("employee_name"))
        leave_label_safe = safe_filename(lt_label).replace("LEAVE", "Leave")
        sponsor_folder_label = safe_filename(sponsor_cfg["heading"])

        emp_dir = outdir / sponsor_cfg["folder"] / f"{emp_name_safe}_{emp.get('sponsor_id')}"
        emp_dir.mkdir(parents=True, exist_ok=True)
        workdir = emp_dir / "_work"
        workdir.mkdir(exist_ok=True)

        log(f"Row {emp.row_num}: generating documents for {emp.get('employee_name')} "
            f"({sponsor_cfg['heading']}, {lt_label})...")
        try:
            if float(str(emp.get("num_days") or 0)) > 90 and not emp.get("visa_days"):
                log(f"    Note: {emp.get('num_days')} leave days exceeds 90 — visa was still "
                    f"marked 90 DAYS. Consider a Multiple-entry visa for this employee.")
        except (TypeError, ValueError):
            pass

        try:
            app_pdf = fill_leave_application(emp, sponsor_name, workdir)
            final_app = emp_dir / f"{emp_name_safe}_{sponsor_folder_label}_{leave_label_safe}_Application.pdf"
            shutil.copy(app_pdf, final_app)

            clr_pdf = fill_clearance_form(emp, sponsor_name, lt_label, workdir)
            final_clr = emp_dir / f"{emp_name_safe}_{sponsor_folder_label}_Clearance.pdf"
            shutil.copy(clr_pdf, final_clr)

            extra_files = []
            if "mehan_vacation" in sponsor_cfg["extra_docs"]:
                mv_pdf = fill_mehan_vacation(emp, lt_key, lt_label, workdir)
                final_mv = emp_dir / f"{emp_name_safe}_{sponsor_folder_label}_VacationRequestFormat.pdf"
                shutil.copy(mv_pdf, final_mv)
                extra_files.append(final_mv)

            if "alfalak_pdf" in sponsor_cfg["extra_docs"]:
                af_pdf = fill_alfalak_pdf(emp, lt_key, workdir)
                final_af = emp_dir / f"{emp_name_safe}_{sponsor_folder_label}_LeaveVisaRequestForm.pdf"
                shutil.copy(af_pdf, final_af)
                extra_files.append(final_af)

            generated.append({
                "row": emp.row_num,
                "name": emp.get("employee_name"),
                "sponsor": sponsor_cfg["heading"],
                "leave_type": lt_label,
                "files": [final_app, final_clr] + extra_files,
            })
        except Exception as exc:
            error_report.append(
                f"Row {emp.row_num} - {emp.get('employee_name') or '(no name)'}\n"
                f"    Generation error: {exc}"
            )
        finally:
            shutil.rmtree(workdir, ignore_errors=True)

    # ---- Summary ----
    log("=" * 70)
    log(f"Processed {len(employees)} employee row(s).")
    log(f"Successfully generated documents for {len(generated)} employee(s).")
    for g in generated:
        log(f"  - Row {g['row']}: {g['name']} [{g['sponsor']} / {g['leave_type']}]")
        for f in g["files"]:
            log(f"      {f}")

    if error_report:
        log("-" * 70)
        log(f"{len(error_report)} issue(s) found — documents NOT generated for these:")
        for e in error_report:
            log(e)
        report_path = outdir / "Error_Report.txt"
        report_path.write_text(
            "Leave Document Generation — Error Report\n"
            f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}\n\n"
            + "\n\n".join(error_report)
        )
        log(f"\nFull error report saved to: {report_path}")
    log("=" * 70)

    return generated, error_report


def main():
    parser = argparse.ArgumentParser(description="Generate leave documents from an Excel input sheet.")
    parser.add_argument("input_excel", help="Path to the filled-in employee Excel sheet")
    parser.add_argument("--outdir", default="Leave Documents", help="Output folder name")
    args = parser.parse_args()
    run_generation(args.input_excel, args.outdir, log=print)


if __name__ == "__main__":
    main()
